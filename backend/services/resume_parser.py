import base64
import io
import os

def get_full_text(file_b64: str, filename: str = "resume.pdf"):

    """Extract full string text from base64 file."""
    file_bytes = base64.b64decode(file_b64)
    ext = os.path.splitext(filename)[-1].lower()

    if ext == ".pdf":
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(file_bytes))
        
        pages_text = []
        for page in reader.pages:
            page_text = page.extract_text() or ""
            links = []
            if "/Annots" in page:
                for annot_ref in page["/Annots"]:
                    try:
                        annot = annot_ref.get_object()
                        if annot.get("/Subtype") == "/Link":
                            a = annot.get("/A")
                            if a:
                                a_obj = a.get_object()
                                if "/URI" in a_obj:
                                    links.append(a_obj["/URI"])
                    except Exception:
                        pass
            if links:
                page_text += "\nLinks: " + ", ".join(links)
            pages_text.append(page_text)
            
        text = "\n".join(pages_text)
        
    elif ext == ".docx":
        import docx
        doc = docx.Document(io.BytesIO(file_bytes))
        text = "\n".join(para.text for para in doc.paragraphs)
        links = []
        for rel in doc.part.rels.values():
            if "hyperlink" in rel.reltype:
                links.append(rel.target_ref)
        if links:
            text += "\nLinks: " + ", ".join(links)
    else:
        text = file_bytes.decode("utf-8", errors="ignore")
    
    return text.strip()

def get_resume_text(file_b64: str, filename: str = "resume.pdf"):

    """
    Decode base64 file and extract text.
    Supports: pdf, txt, docx, doc
    Returns: list of text chunks (plain strings)
    """
    file_bytes = base64.b64decode(file_b64)
    ext = os.path.splitext(filename)[-1].lower()

    # ── Extract raw text based on file type ──────────────────────
    if ext == ".pdf":
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(file_bytes))
        
        pages_text = []
        for page in reader.pages:
            page_text = page.extract_text() or ""
            links = []
            if "/Annots" in page:
                for annot_ref in page["/Annots"]:
                    try:
                        annot = annot_ref.get_object()
                        if annot.get("/Subtype") == "/Link":
                            a = annot.get("/A")
                            if a:
                                a_obj = a.get_object()
                                if "/URI" in a_obj:
                                    links.append(a_obj["/URI"])
                    except Exception:
                        pass
            if links:
                page_text += "\nLinks: " + ", ".join(links)
            pages_text.append(page_text)
            
        text = "\n".join(pages_text)

    elif ext == ".docx":
        import docx
        doc = docx.Document(io.BytesIO(file_bytes))
        text = "\n".join(para.text for para in doc.paragraphs)
        links = []
        for rel in doc.part.rels.values():
            if "hyperlink" in rel.reltype:
                links.append(rel.target_ref)
        if links:
            text += "\nLinks: " + ", ".join(links)

    elif ext == ".txt":
        text = file_bytes.decode("utf-8", errors="ignore")

    else:
        # try plain text as fallback
        text = file_bytes.decode("utf-8", errors="ignore")

    text = text.strip()
    if not text:
        return []

    # ── Simple chunking (800 chars, 150 overlap) ──────────────────
    chunk_size = 800
    overlap    = 150
    chunks = []
    start  = 0
    while start < len(text):
        chunks.append(text[start : start + chunk_size])
        start += chunk_size - overlap

    return chunks
